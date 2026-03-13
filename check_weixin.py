# -*- coding: utf-8 -*-
"""
检查Word文档中的微信公众号链接
"""
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re

doc_path = r'C:\Users\Administrator\Desktop\Shiyo\数据分析部\游戏行业知识库\游戏数据分析知识库\Int 副本.docx'
doc = Document(doc_path)

# 方法1: 检查超链接
hyperlink_weixin = []
for rel in doc.part.rels.values():
    if rel.reltype == RT.HYPERLINK:
        url = rel.target_ref
        if 'weixin.qq.com' in url or 'mp.weixin' in url:
            hyperlink_weixin.append(url)

# 方法2: 检查纯文本内容
all_text = []
for para in doc.paragraphs:
    all_text.append(para.text)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_text.append(cell.text)

full_text = '\n'.join(all_text)

# 搜索微信公众号链接
weixin_pattern = r'https?://mp\.weixin\.qq\.com/[^\s<>"\)]+'
text_weixin_links = re.findall(weixin_pattern, full_text)

# 输出结果
with open('weixin_check_result.txt', 'w', encoding='utf-8') as f:
    f.write(f'=== Hyperlink check ===\n')
    f.write(f'Found {len(hyperlink_weixin)} weixin links in hyperlinks\n')
    for link in hyperlink_weixin:
        f.write(f'  {link}\n')
    
    f.write(f'\n=== Text content check ===\n')
    f.write(f'Found {len(text_weixin_links)} weixin links in text\n')
    for link in text_weixin_links:
        f.write(f'  {link}\n')
    
    f.write(f'\n=== All unique weixin links ===\n')
    all_weixin = set(hyperlink_weixin + text_weixin_links)
    f.write(f'Total unique: {len(all_weixin)}\n')
    for i, link in enumerate(all_weixin, 1):
        f.write(f'{i}. {link}\n')

print('Result saved to weixin_check_result.txt')
